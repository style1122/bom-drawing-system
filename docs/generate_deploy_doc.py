# -*- coding: utf-8 -*-
"""生成 BOM 图纸管理系统 Windows Server 2019 部署文档 (.docx)"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"E:\workspace\bom-drawing-system\docs\BOM图纸管理系统_部署文档_WindowsServer2019.docx"

CJK = "Microsoft YaHei"
MONO = "Consolas"

doc = Document()

# ---- 默认字体（含中文） ----
def set_cjk(style_or_run, font=CJK):
    if hasattr(style_or_run, "font"):
        r = style_or_run.font
    else:
        r = style_or_run
    r.name = font
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)

normal = doc.styles['Normal']
normal.font.size = Pt(10.5)
set_cjk(normal)

# 标题配色
TITLE_COLOR = RGBColor(0x1F, 0x4E, 0x79)
H_COLOR = RGBColor(0x2E, 0x5C, 0x8A)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = TITLE_COLOR
    set_cjk(r)
    return p

def add_heading(text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_cjk(r)
    r.font.color.rgb = H_COLOR
    return p

def add_para(text, bold=False, size=10.5, color=None, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    set_cjk(r)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if level:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    r = p.add_run(text)
    set_cjk(r)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = MONO
        r.font.size = Pt(9.5)
        r._element.get_or_add_rPr().append(OxmlElement('w:rFonts'))
        r._element.rPr.find(qn('w:rFonts')).set(qn('w:ascii'), MONO)
        r._element.rPr.find(qn('w:rFonts')).set(qn('w:hAnsi'), MONO)
    # 浅底纹
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0].add_run(h)
        rp.bold = True
        set_cjk(rp)
        rp.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rp = cells[i].paragraphs[0].add_run(str(val))
            set_cjk(rp)
            rp.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

# =====================================================================
add_title("BOM 图纸管理系统")
add_para("Windows Server 2019 生产环境部署文档", bold=True, size=13, color=H_COLOR).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_para("文档版本：v1.1    生成日期：2026-09-02    适用系统：BOM 图纸管理系统（Spring Boot + MyBatis-Plus 版）",
         size=9, color=RGBColor(0x66,0x66,0x66)).alignment = WD_ALIGN_PARAGRAPH.CENTER

# 0. 文档信息
add_heading("0. 文档信息", 1)
add_table(
    ["项目", "说明"],
    [
        ["系统名称", "BOM 图纸管理系统"],
        ["后端技术栈", "Spring Boot 2.7.18（内嵌 Tomcat 9 / javax.servlet）+ MyBatis-Plus 3.5.7 + Druid"],
        ["数据库", "SQL Server 2019（混合验证模式）"],
        ["前端技术栈", "Vue 3 + Vite，构建为静态资源由 nginx 托管"],
        ["部署形态", "后端：可执行 JAR（单文件，无需外置 Tomcat）；前端：静态文件 + nginx 反向代理"],
        ["操作系统", "Windows Server 2019（64 位）"],
        ["代码分支", "main（GitHub：git@github.com:style1122/bom-drawing-system.git）"],
    ],
    widths=[1.6, 4.8],
)

# 1. 系统概述与部署架构
add_heading("1. 系统概述与部署架构", 1)
add_para("本系统用于管理 BOM 图纸与物料基础数据，后端提供 /api/** 接口，前端为单页应用（SPA）。"
         "生产环境采用「nginx 托管前端 + 反向代理后端」的拓扑，后端以 Windows 服务形式常驻运行。")
add_para("部署拓扑：")
add_code(
"[ 浏览器 ]  --HTTP 80-->  [ nginx for Windows ]\n"
"                              |  /             --> 静态文件：frontend/dist\n"
"                              |  /api/**       --> 反向代理 --> [ Spring Boot JAR :8080 ]\n"
"                                                                     |\n"
"                                                                     --> [ SQL Server 2019 :1433 / BOM_DB ]"
)
add_bullet("后端只提供 /api/** 接口；前端页面与静态资源由 nginx 直接返回。")
add_bullet("后端注册为 Windows 服务（WinSW），开机自启、崩溃自动重启。")
add_bullet("所有密钥（数据库密码、ERP 密钥）通过环境变量注入，不落明文。")

# 2. 环境要求
add_heading("2. 环境要求", 1)
add_table(
    ["组件", "版本", "说明"],
    [
        ["Windows Server", "2019（64 位）", "目标部署操作系统"],
        ["JDK", "8uXXX 64-bit（或 11/17）", "Spring Boot 2.7 支持 Java 8–19；需 java 在 PATH"],
        ["SQL Server", "2019", "数据库引擎，混合验证模式"],
        ["Maven", "3.6+", "仅构建机需要（打后端 JAR）"],
        ["Node.js", "18+", "仅构建前端需要（含 npm）"],
        ["nginx", "for Windows 1.24+", "托管前端 + 反向代理"],
        ["WinSW", "v2 / v3 (x64)", "将 JAR 注册为 Windows 服务"],
    ],
    widths=[1.4, 2.0, 3.0],
)
add_para("说明：构建机与部署机可为同一台，也可分开。若分开，只需把构建产物 "
         "bom-drawing-system.jar 与 frontend/dist 拷贝到服务器即可。", italic=True, size=9.5)

# 3. 数据库准备
add_heading("3. 数据库准备（SQL Server 2019）", 1)
add_heading("3.1 安装与网络配置", 2)
add_bullet("安装 SQL Server 2019 数据库引擎，身份验证选【混合模式】，设置 sa 密码（该密码即后续 JDBC_PASSWORD）。")
add_bullet("打开 SQL Server Configuration Manager → SQL Server Network Configuration → Protocols for MSSQLSERVER → TCP/IP → 启用（Yes）。")
add_bullet("同窗口 TCP/IP → IP Addresses → 拉到 IPAll → 设置 TCP Port = 1433（清空 TCP Dynamic Ports）。")
add_bullet("重启 SQL Server (MSSQLSERVER) 服务使配置生效。")
add_bullet("Windows 防火墙 → 高级设置 → 入站规则 → 新建规则 → 端口 → TCP 1433 → 允许。")
add_heading("3.2 建库与建表", 2)
add_para("用 SSMS 或 sqlcmd 依次执行项目 sql/ 目录下的脚本（建议数据库排序规则 Chinese_PRC_CI_AS）：")
add_code(
"sql/01_create_database.sql     # 建库 BOM_DB\n"
"sql/02_create_requisition.sql  # 请购单相关表\n"
"sql/03_erp_material.sql        # ERP 物料基础数据表\n"
"sql/03_update_roles.sql        # 角色初始化\n"
"sql/04_erp_have_drawing.sql    # ERP 是否已出图标记表\n"
"sql/05_erp_sync_cursor.sql     # ERP 同步游标表"
)
add_para("若使用专用应用账号（而非 sa），请在 BOM_DB 中创建登录名并映射为 db_owner（或最小必需权限）。", italic=True, size=9.5)

# 4. 获取源码与构建
add_heading("4. 获取源码与构建", 1)
add_heading("4.1 获取代码", 2)
add_para("在构建机上克隆仓库（已配置 SSH 密钥）：")
add_code("git clone git@github.com:style1122/bom-drawing-system.git\ncd bom-drawing-system")
add_heading("4.2 构建", 2)
add_para("仓库根目录已提供 build.bat，一键完成「后端 Maven 打包 + 前端 npm 构建」：")
add_code("build.bat")
add_para("或手动执行：")
add_code(
"mvn clean package -DskipTests\n"
"cd frontend && npm install && npm run build"
)
add_para("构建产物：", bold=True)
add_bullet("后端：backend\\target\\bom-drawing-system.jar（finalName=bom-drawing-system）")
add_bullet("前端：frontend\\dist\\ （静态资源）")
add_para("注意：若构建机无法联网执行 mvn，请在可联网的构建机上完成打包后再拷贝 jar 到服务器。",
         italic=True, size=9.5)

# 5. 配置密钥
add_heading("5. 配置密钥（环境变量）", 1)
add_para("application.yml 中数据库与 ERP 密钥均使用 ${ENV_VAR} 占位并保留本地默认值。"
         "生产服务器必须通过系统环境变量覆盖敏感项（尤其 JDBC_PASSWORD、ERP_APPSECRET）。")
add_para("在服务器上设置【系统环境变量】（控制面板 → 系统 → 高级 → 环境变量 → 系统变量）：")
add_table(
    ["变量名", "说明", "默认值（开发）"],
    [
        ["JDBC_URL", "SQL Server 连接串", "jdbc:sqlserver://localhost:1433;databaseName=BOM_DB;encrypt=false;trustServerCertificate=true;sendStringParametersAsUnicode=true"],
        ["JDBC_USERNAME", "数据库账号", "sa"],
        ["JDBC_PASSWORD", "数据库密码", "Sync@2026（生产务必修改）"],
        ["SERVER_PORT", "后端监听端口", "8080"],
        ["ERP_BASE_URL", "正航 ESB 地址", "http://10.1.1.15:860"],
        ["ERP_APPID", "ERP 应用 ID", "见 application.yml"],
        ["ERP_APPSECRET", "ERP 密钥", "见 application.yml（生产务必修改）"],
    ],
    widths=[1.5, 1.7, 3.2],
)
add_para("WinSW 启动的 Java 进程会继承系统环境变量，因此配置在系统变量中即可被服务读取。", italic=True, size=9.5)

# 6. 部署后端
add_heading("6. 部署后端（JAR + WinSW 服务）", 1)
add_bullet("在服务器建目录 C:\\apps\\bom-drawing-system\\。")
add_bullet("拷贝：bom-drawing-system.jar、deploy\\windows\\bom-drawing-system.xml。")
add_bullet("从 https://github.com/winsw/winsw/releases 下载 WinSW-x64.exe，重命名为 bom-drawing-system.exe，与本目录其余文件同级。")
add_para("以管理员身份运行安装服务：")
add_code(
"cd C:\\apps\\bom-drawing-system\n"
"install-service.bat"
)
add_para("启动服务（或在 services.msc 中找到 “BOM Drawing System”）：")
add_code("bom-drawing-system.exe start")
add_para("常用命令：")
add_code(
"bom-drawing-system.exe stop      # 停止\n"
"bom-drawing-system.exe restart   # 重启\n"
"uninstall-service.bat            # 卸载服务"
)
add_para("验证：", bold=True)
add_bullet("日志：C:\\apps\\bom-drawing-system\\logs\\bom-system.log（应用）与 bom-drawing-system.out.log（WinSW 控制台）。")
add_bullet("接口探活：curl http://localhost:8080/api/...（登录类接口已放行）。")

# 7. 部署前端
add_heading("7. 部署前端（nginx）", 1)
add_bullet("建目录 C:\\apps\\bom-drawing-system\\frontend\\，将 frontend\\dist\\ 整个拷贝到 ...\\frontend\\dist。")
add_bullet("安装 nginx for Windows，将 deploy\\nginx\\bom-drawing-system.conf 复制到 nginx\\conf\\。")
add_bullet("在 nginx\\conf\\nginx.conf 的 http { ... } 内添加：include bom-drawing-system.conf;")
add_para("以管理员身份启动 nginx：")
add_code(
"cd C:\\apps\\nginx\n"
"start nginx.exe"
)
add_para("验证：浏览器访问 http://<服务器IP>/ 应打开登录页；接口请求走 /api/** 被代理到后端。")
add_para("若 80 端口被 IIS 占用，可改 listen 80; 为其他端口（如 8081），或停用 IIS 默认站点。",
         italic=True, size=9.5)

# 8. 目录结构
add_heading("8. 部署后目录结构", 1)
add_code(
"C:\\apps\\\n"
"+-- bom-drawing-system\\\n"
"|   +-- bom-drawing-system.jar\n"
"|   +-- bom-drawing-system.exe        (WinSW)\n"
"|   +-- bom-drawing-system.xml        (WinSW 配置)\n"
"|   +-- logs\\                          (应用/服务日志)\n"
"|   +-- frontend\\\n"
"|       +-- dist\\                      (Vue 构建产物)\n"
"+-- nginx\\\n"
"    +-- conf\\\n"
"        +-- nginx.conf\n"
"        +-- bom-drawing-system.conf"
)

# 9. 运维与故障排查
add_heading("9. 运维与故障排查", 1)
add_heading("9.1 日志", 2)
add_bullet("应用日志：C:\\apps\\bom-drawing-system\\logs\\bom-system.log（按天滚动，保留 30 天）。")
add_bullet("服务控制台：bom-drawing-system.out.log / .err.log（WinSW 生成）。")
add_bullet("nginx 日志：nginx\\logs\\access.log / error.log。")
add_heading("9.2 端口冲突", 2)
add_bullet("若 8080 被占用：设置系统环境变量 SERVER_PORT=8090（或改 application.yml 的 server.port），并同步修改 nginx proxy_pass 端口。")
add_bullet("nginx 80 被占用：见第 7 步（改用其他端口或停用 IIS）。")
add_heading("9.3 常见错误", 2)
add_table(
    ["现象", "原因 / 处理"],
    [
        ["启动报 datetime2 / 无效的列类型", "必须保留 mybatis-plus.configuration.jdbc-type-for-null: NULL（已在 application.yml 配置），确保未误删。"],
        ["连接 SQL Server 报 TLS / 加密错误", "开发用 encrypt=false;trustServerCertificate=true；生产启用 TLS 时改 encrypt=true 并配置证书。"],
        ["API 返回 401", "未携带/已失效 token；登录、注册、公开分享、下载、ERP 订阅接口已放行。"],
        ["API 返回 403", "写操作需 ADMIN/ENGINEER 角色（见 RoleInterceptor 映射）。"],
        ["服务启动即退出", "查 out.log：多为 java 未在 PATH、或 JDBC_PASSWORD 等系统环境变量未设置。"],
        ["前端页面白屏/刷新 404", "确认 nginx location / 有 try_files $uri $uri/ /index.html;（SPA 路由兜底）。"],
    ],
    widths=[2.2, 4.2],
)
add_heading("9.4 升级流程", 2)
add_bullet("构建机重新 build.bat 生成新 jar / dist。")
add_bullet("停止服务：bom-drawing-system.exe stop。")
add_bullet("替换 bom-drawing-system.jar 与 frontend\\dist。")
add_bullet("启动服务：bom-drawing-system.exe start；如改了 nginx 配置需 nginx -s reload。")

# 10. 安全与最佳实践
add_heading("10. 安全与最佳实践", 1)
add_bullet("密钥一律通过环境变量注入，禁止将 JDBC_PASSWORD / ERP_APPSECRET 明文写入 application.yml 或提交到仓库。")
add_bullet("生产数据库账号建议使用专用低权限账号，而非 sa。")
add_bullet("Windows 防火墙仅开放必要端口（80/443 对外，1433 建议仅限内网或绑定内网网卡）。")
add_bullet("nginx 建议启用 HTTPS（listen 443 ssl）并对 /api 做必要的限流与鉴权加固。")
add_bullet("定期备份 BOM_DB 数据库与 C:\\apps\\bom-drawing-system\\logs 日志。")

# 附录
add_heading("附录 A：application.yml 关键配置", 1)
add_bullet("server.port = ${SERVER_PORT:8080}，context-path = /")
add_bullet("文件上传上限 500MB（与原有配置一致），大文件上传已对应延长 nginx 超时。")
add_bullet("Druid 连接池：initial-size 5 / min-idle 5 / max-active 20。")
add_bullet("MyBatis-Plus：jdbc-type-for-null=NULL（关键，避免 datetime2 写入报错）；map-underscore-to-camel-case=true。")
add_bullet("ERP 对接：erp.base-url / erp.appid / erp.appsecret 均支持环境变量覆盖；订阅轮询默认开启（interval-ms=300000）。")

add_heading("附录 B：正航 T9 ERP 对接要点", 1)
add_bullet("物料基础数据通过正航 ESB（Mat01）拉取，OAuth2 client_credentials + ESB MD5 签名。")
add_bullet("数据库需存在 erp_material、erp_sync_cursor 等表（由 sql/03、05 脚本创建）。")
add_bullet("同步游标表记录上次同步时间戳，断点续传；首次可从 2018-01-01 起全量。")

add_para("")
add_para("— 文档结束 —", italic=True, size=9, color=RGBColor(0x99,0x99,0x99)).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT)
print("SAVED:", OUT)
