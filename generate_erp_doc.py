# -*- coding: utf-8 -*-
"""
生成正航T9 ERP对接需求文档
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== 辅助函数 =====
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return table

# ===== 封面 =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('\n\n\nBOM图纸管理系统\n与正航T9 ERP对接\n需求资料清单')
run.font.size = Pt(28)
run.bold = True
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.color.rgb = RGBColor(0x1a, 0x5c, 0xb0)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('\n基于正航ERP API接口文档V2.2\n2026年7月')
run.font.size = Pt(14)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ===== 目录 =====
add_heading('目 录', level=1)
toc_items = [
    '一、文档概述',
    '二、正航T9 ERP API 认证机制',
    '三、需要从正航获取的资料清单',
    '四、系统功能与ERP接口对照总表',
    '五、接口地址清单',
    '六、数据字段映射（物料同步）',
    '七、数据字段映射（采购收货单同步）',
    '八、数据字段映射（BOM同步 — 待实现）',
    '九、数据字段映射（生产订单同步 — 待实现）',
    '十、辅助接口（人员/部门/仓库）',
    '十一、系统现状与缺口分析',
    '十二、实施建议与下一步',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ===== 一、文档概述 =====
add_heading('一、文档概述', level=1)
add_para('本文档基于《正航ERP API接口文档V2.2》，结合BOM图纸管理系统现有功能模块，梳理系统与正航T9 ERP对接所需的全部资料，包括：')
add_para('1. 从正航获取的认证凭据（appid / appsecret / ESB服务地址）')
add_para('2. 需要调用的ERP接口地址及功能编号（progid）')
add_para('3. 每个对接功能的数据字段映射关系（系统表字段 ↔ ERP API字段）')
add_para('4. 系统现有功能与ERP接口的差距分析')
add_para('5. 实施建议和下一步行动项')

add_para('')
add_para('系统现有功能模块：', bold=True)
add_table(
    ['模块', '数据库表', '当前数据来源', '对接ERP接口'],
    [
        ['物料管理', 'material', '手动录入', 'comMaterialGroup（物料基础数据）'],
        ['图纸管理', 'drawing', '本地上传', '无对应ERP接口（系统独有）'],
        ['采购订单管理', 'requisition + requisition_item', '手动录入/测试数据', 'purReceivingOrder_RC（采购收货单）'],
        ['用户管理', 'sys_user', '注册', '无对应ERP接口'],
        ['BOM管理', '无（未实现）', '—', 'ppBOM（BOM）'],
        ['生产订单', '无（未实现）', '—', 'ppProduceOrder（生产订单）'],
    ]
)

doc.add_page_break()

# ===== 二、认证机制 =====
add_heading('二、正航T9 ERP API 认证机制', level=1)

add_heading('2.1 认证流程', level=2)
add_para('正航ERP采用ESB（企业服务总线）架构，所有业务API调用前必须先通过认证接口获取token。')
add_para('')
add_para('流程：', bold=True)
add_para('① 系统调用认证接口 → ② ESB验证appid+appsecret → ③ 返回token（有效期7200秒/2小时）')
add_para('④ 后续业务API请求携带token → ⑤ token过期后需重新认证')

add_heading('2.2 认证接口', level=2)
add_table(
    ['项目', '值'],
    [
        ['接口地址', 'http://(ESB服务地址及端口)/esb/api/auth.do'],
        ['请求方式', 'HTTP POST'],
        ['请求格式', 'JSON (Content-Type: application/json)'],
        ['响应格式', 'JSON'],
        ['编码格式', 'UTF-8'],
    ]
)

add_heading('2.3 认证请求参数', level=2)
add_table(
    ['参数名', '类型', '必填', '说明'],
    [
        ['appid', 'String', 'Y', '应用ID，由正航提供'],
        ['time', 'String/Long', 'Y', 'UTC时间戳（毫秒）+ 6位随机数'],
        ['sign', 'String', 'Y', '签名 = MD5(appid + appsecret + time)，大写'],
        ['language', 'String', 'N', '语言：zh-CHS(简体)/zh-CHT(繁體)/en'],
    ]
)

add_heading('2.4 签名生成规则', level=2)
add_para('sign = MD5(appid + appsecret + time)，结果转大写')
add_para('')
add_para('示例：', bold=True)
add_para('appid = chiesb185a1fbdc22781f1')
add_para('appsecret = 2BB2D1D899C32BF091BEC8DDA9EBCB72')
add_para('time(UTC毫秒) = 1718855090958')
add_para('time(+6位随机数) = 1718855090958191519')
add_para('拼接字符串 = chiesb185a1fbdc22781f12BB2D1D899C32BF091BEC8DDA9EBCB721718855090958191519')
add_para('MD5结果(sign) = CFCCF02D2D1370757F295472A2F597D1（大写）')

add_heading('2.5 业务API签名规则（与认证不同）', level=2)
add_para('业务API（getlist/getdata/addnew等）的sign生成方式与认证接口不同：')
add_para('sign = MD5(token + appsecret + time)，结果转大写')
add_para('注意：业务API的time只需要UTC毫秒时间戳 + 1位随机数（认证接口是6位）')

add_heading('2.6 认证成功响应', level=2)
add_table(
    ['参数名', '类型', '说明'],
    [
        ['status', 'Int', '1=成功'],
        ['token', 'String', '身份令牌（32位），后续API调用使用'],
        ['timeout', 'Long', 'token有效期（秒），通常7200（2小时）'],
        ['time', 'Long', '服务端UTC时间'],
        ['sign', 'String', '响应签名 = MD5(appid+appsecret+token+time)'],
    ]
)

doc.add_page_break()

# ===== 三、资料清单 =====
add_heading('三、需要从正航获取的资料清单', level=1)

add_para('以下资料必须从正航软件方或企业ERP管理员处获取，否则无法启动对接开发：', bold=True, color=(0xcc, 0x00, 0x00))

add_heading('3.1 必须获取的凭据（⚠️ 阻塞项）', level=2)
add_table(
    ['序号', '资料名称', '说明', '示例值（文档中的测试值）', '获取来源'],
    [
        ['1', 'ESB服务地址及端口', '所有API调用的基础URL', 'http://192.168.x.x:8080', '正航实施人员/IT部门'],
        ['2', 'appid', '应用唯一标识', 'chiesb185a1fbdc22781f1', '正航提供'],
        ['3', 'appsecret', '应用密钥（用于签名）', '2BB2D1D899C32BF091BEC8DDA9EBCB72', '正航提供'],
    ]
)

add_heading('3.2 需要确认的基础数据代码对照表', level=2)
add_para('ERP中的物料类型、物料类别、计量单位、部门、人员等均使用代码，需要获取ERP中实际使用的代码值清单：')
add_table(
    ['序号', '数据项', 'ERP接口', '用途', '需要的内容'],
    [
        ['1', '物料类型代码', 'comMaterialGroup.MaterialTypeId', '物料分类', '如：01原材料/02半成品/03成品 等ERP实际代码'],
        ['2', '物料类别代码', 'comMaterialGroup.MaterialCategoryId', '物料分类（另一维度）', 'ERP后台物料类别清单'],
        ['3', '计量单位代码', 'comMaterialGroup.UnitId', '物料的计量单位', '如：PCS/SET/KG 等ERP实际代码'],
        ['4', 'BOM类型代码', 'ppBOM.BOMStyleId', 'BOM类型', '文档示例固定值0001，需确认'],
        ['5', 'BOM用途类别', 'ppBOM.BOMTypeId', 'BOM用途', 'ERP后台BOM用途类别清单'],
        ['6', '单据类型代码', 'purReceivingOrder_RC.TypeId', '采购收货单类型', 'ERP中实际使用的单据类型代码'],
        ['7', '人员代码对照', 'comCompanyPerson', '采购人员PersonId ↔ 姓名', 'ERP人员代码与姓名的对照表'],
        ['8', '部门代码对照', 'comDepartment', '部门DeptId ↔ 部门名称', 'ERP部门代码与名称的对照表'],
    ]
)

add_heading('3.3 需要确认的业务规则', level=2)
add_table(
    ['序号', '确认项', '说明', '影响'],
    [
        ['1', '物料编码规则', 'ERP中MaterialId的编码规则和格式', '系统material表的material_code需与ERP一致才能关联图纸'],
        ['2', '采购人员匹配方式', 'ERP的PersonId(代码) 与 系统的display_name(姓名) 如何对应', '影响采购角色数据过滤功能'],
        ['3', '采购订单对应关系', '系统的"采购订单"对应ERP的哪种单据？文档中仅有purReceivingOrder_RC(采购收货单)', '需确认ERP中是否有请购单或采购订单接口'],
        ['4', '同步频率', '物料/采购单/BOM的同步频率要求', '实时/定时（如每小时/每天）'],
        ['5', '数据范围', '需要同步全量还是增量数据', '影响getlist的condition条件和分页策略'],
        ['6', '公司代码', 'comMaterial需要FOrgId(公司代码)作为查询条件', '需要确认ERP中的公司组织代码'],
    ]
)

doc.add_page_break()

# ===== 四、功能对照总表 =====
add_heading('四、系统功能与ERP接口对照总表', level=1)

add_heading('4.1 ERP API完整业务对象列表', level=2)
add_para('正航T9 ERP API提供以下业务对象接口，每个对象支持：新增/修改/删除/清单查询/明细查询/审核/作废/结案')
add_table(
    ['序号', '业务对象', 'progid', '中文名称', '与本系统相关性'],
    [
        ['1', 'comCustomer', 'comCustomer', '客户资料', '低'],
        ['2', 'comSupplier', 'comSupplier', '供应商', '低'],
        ['3', 'comMaterialGroup', 'comMaterialGroup', '物料基础数据', '★ 高（物料同步）'],
        ['4', 'comMaterial', 'comMaterial', '物料公司数据', '★ 高（物料同步）'],
        ['5', 'comDepartment', 'comDepartment', '部门', '中（部门对照）'],
        ['6', 'stkWareHouse', 'stkWareHouse', '仓库', '低'],
        ['7', 'comGroupPerson', 'comGroupPerson', '人员集团档案', '中（人员对照）'],
        ['8', 'comCompanyPerson', 'comCompanyPerson', '人员人事档案', '★ 高（采购人员匹配）'],
        ['9', 'plsTechRoute', 'plsTechRoute', '工艺路线', '低'],
        ['10', 'plsWorkCenter', 'plsWorkCenter', '工作中心', '低'],
        ['11', 'plsWorkingProcedureCode', 'plsWorkingProcedureCode', '工序代码', '低'],
        ['12', 'ppBOM', 'ppBOM', 'BOM', '★ 高（BOM同步 — 待实现）'],
        ['13', 'ppBOMBatchVaryBill', 'ppBOMBatchVaryBill', 'BOM子件批次维护单', '低'],
        ['14', 'ppProduceOrder', 'ppProduceOrder', '生产订单', '中（生产订单同步 — 待实现）'],
        ['15', 'ppProduceSend', 'ppProduceSend', '生产发料', '低'],
        ['16', 'ppTakeMatApply_MR', 'ppTakeMatApply_MR', '领料申请单', '低'],
        ['17', 'ppWorkReport', 'ppWorkReport', '工报单', '低'],
        ['18', 'ppProduceStoreIn_PO', 'ppProduceStoreIn_PO', '生产入库申请单', '低'],
        ['19', 'salSalesOrder_SO', 'salSalesOrder_SO', '销售订单', '低'],
        ['20', 'purReceivingOrder_RC', 'purReceivingOrder_RC', '采购收货单', '★ 高（采购订单同步）'],
        ['21', 'stkOtherStockOut_PM', 'stkOtherStockOut_PM', '其他出库申请单', '低'],
        ['22', 'stkOtherStockIn_OI', 'stkOtherStockIn_OI', '其他入库申请单', '低'],
        ['23', 'salDispatchList_DL', 'salDispatchList_DL', '销售发货单', '低'],
        ['24', 'salDispatchList_GR', 'salDispatchList_GR', '销售发货单', '低'],
    ]
)

add_heading('4.2 本系统需要的ERP接口（按优先级）', level=2)
add_table(
    ['优先级', '系统功能', 'ERP接口progid', '操作', '用途'],
    [
        ['P0 必须', '物料同步', 'comMaterialGroup', 'getlist', '从ERP拉取物料基础数据到material表'],
        ['P0 必须', '认证', 'auth', '认证', '获取token，所有API调用的前提'],
        ['P1 重要', '采购订单同步', 'purReceivingOrder_RC', 'getlist + getdata', '从ERP拉取采购收货单到requisition表'],
        ['P1 重要', '人员对照', 'comCompanyPerson', 'getlist', '获取人员代码↔姓名对照，用于采购人员匹配'],
        ['P2 可选', 'BOM同步', 'ppBOM', 'getlist + getdata', '从ERP拉取BOM数据（系统尚未实现BOM功能）'],
        ['P2 可选', '部门对照', 'comDepartment', 'getlist', '获取部门代码↔名称对照'],
        ['P3 远期', '生产订单同步', 'ppProduceOrder', 'getlist', '从ERP拉取生产订单数据'],
    ]
)

doc.add_page_break()

# ===== 五、接口地址清单 =====
add_heading('五、接口地址清单', level=1)

add_para('所有接口的基础URL = http://(ESB服务地址及端口)')
add_para('')

add_heading('5.1 认证接口', level=2)
add_table(
    ['功能', 'HTTP方法', '完整路径', '说明'],
    [
        ['获取token', 'POST', '/esb/api/auth.do', '使用appid+appsecret认证获取token'],
    ]
)

add_heading('5.2 通用业务接口', level=2)
add_table(
    ['功能', 'HTTP方法', '完整路径', '说明'],
    [
        ['新增数据', 'POST', '/esb/erp/addnew.do', '新增业务单据/基础数据'],
        ['修改数据', 'POST', '/esb/erp/update.do', '修改业务单据/基础数据'],
        ['删除数据', 'POST', '/esb/erp/delete.do', '删除业务单据/基础数据'],
        ['清单查询', 'POST', '/esb/erp/getlist.do', '分页查询列表数据'],
        ['明细查询', 'POST', '/esb/erp/get.do', '查询单据明细（含子表）'],
        ['审核', 'POST', '/esb/erp/approve.do', '审核单据'],
        ['作废', 'POST', '/esb/erp/invalidate.do', '作废单据'],
        ['结案', 'POST', '/esb/erp/endcase.do', '结案单据'],
    ]
)

add_heading('5.3 本系统需要调用的接口清单', level=2)
add_table(
    ['序号', '用途', '接口路径', 'progid', '调用频率'],
    [
        ['1', '获取认证token', '/esb/api/auth.do', '—', '每2小时'],
        ['2', '查询物料基础数据', '/esb/erp/getlist.do', 'comMaterialGroup', '每天/手动触发'],
        ['3', '查询物料公司数据', '/esb/erp/getlist.do', 'comMaterial', '每天/手动触发'],
        ['4', '查询采购收货单列表', '/esb/erp/getlist.do', 'purReceivingOrder_RC', '每小时/手动触发'],
        ['5', '查询采购收货单明细', '/esb/erp/get.do', 'purReceivingOrder_RC', '按需（getlist后逐单获取）'],
        ['6', '查询人员档案', '/esb/erp/getlist.do', 'comCompanyPerson', '每天'],
        ['7', '查询部门', '/esb/erp/getlist.do', 'comDepartment', '每周'],
        ['8', '查询BOM列表', '/esb/erp/getlist.do', 'ppBOM', '每天/手动触发'],
        ['9', '查询BOM明细', '/esb/erp/get.do', 'ppBOM', '按需（getlist后逐单获取）'],
        ['10', '查询生产订单', '/esb/erp/getlist.do', 'ppProduceOrder', '每小时'],
    ]
)

add_heading('5.4 getlist通用请求参数', level=2)
add_table(
    ['参数名', '类型', '必填', '说明'],
    [
        ['token', 'String', 'Y', '认证获取的token'],
        ['time', 'String/Long', 'Y', 'UTC时间戳(毫秒) + 1位随机数'],
        ['sign', 'String', 'Y', 'MD5(token + appsecret + time)，大写'],
        ['progid', 'String', 'Y', '功能编号（如comMaterialGroup）'],
        ['data.condition', 'String', 'N', 'SQL Where条件（如 "MaterialId LIKE \'A%\'"）'],
        ['data.lastpkvalues', 'String', 'N', '分页：上一页最后一条的主键值'],
    ]
)

add_heading('5.5 getdata通用请求参数', level=2)
add_table(
    ['参数名', '类型', '必填', '说明'],
    [
        ['token', 'String', 'Y', '认证获取的token'],
        ['time', 'String/Long', 'Y', 'UTC时间戳(毫秒) + 1位随机数'],
        ['sign', 'String', 'Y', 'MD5(token + appsecret + time)，大写'],
        ['progid', 'String', 'Y', '功能编号'],
        ['data.pkvalues', 'String', 'Y', '主键值，多主键用逗号分隔'],
    ]
)

doc.add_page_break()

# ===== 六、物料同步字段映射 =====
add_heading('六、数据字段映射（物料同步）', level=1)

add_heading('6.1 接口信息', level=2)
add_table(
    ['项目', '值'],
    [
        ['ERP接口', '/esb/erp/getlist.do'],
        ['progid', 'comMaterialGroup'],
        ['ERP数据表', 'MaterialGroup'],
        ['系统数据表', 'material'],
        ['同步方向', 'ERP → 系统（单向拉取）'],
    ]
)

add_heading('6.2 字段映射表', level=2)
add_table(
    ['系统字段(material表)', '类型', 'ERP字段(MaterialGroup)', 'ERP类型', '映射说明', '备注'],
    [
        ['material_code', 'NVARCHAR(64)', 'MaterialId', 'String(40)', '直接映射', '主键，ERP物料代码'],
        ['material_name', 'NVARCHAR(256)', 'MaterialName', 'String(120)', '直接映射', '物料名称'],
        ['specification', 'NVARCHAR(512)', 'MaterialSpec', 'String(200)', '直接映射', '物料规格'],
        ['material_type', 'NVARCHAR(32)', 'MaterialTypeId', 'String(10)', '代码→中文转换', '需获取代码对照表'],
        ['unit', 'NVARCHAR(16)', 'UnitId', 'String(20)', '代码→中文转换', '需获取单位对照表'],
        ['—', '—', 'MaterialCategoryId', 'String(20)', '系统无对应字段', '物料类别，建议新增字段'],
        ['—', '—', 'ValidityFromDate', 'String(10)', '系统无对应字段', '有效期从，可选同步'],
        ['—', '—', 'ValidityToDate', 'String(10)', '系统无对应字段', '有效期至，可选同步'],
        ['—', '—', 'UseDealMultiUnit', 'Boolean', '系统无对应字段', '使用交易多单位，可不同步'],
        ['drawing_no', 'NVARCHAR(128)', '—', '—', '系统独有', '图纸编号，系统维护'],
        ['weight', 'DECIMAL(18,6)', '—', '—', '系统独有', '重量，系统维护'],
        ['material_attr', 'NVARCHAR(128)', '—', '—', '系统独有', '物料属性，系统维护'],
        ['source', 'NVARCHAR(32)', '—', '—', '系统独有', '数据来源，同步时设为ERP'],
        ['erp_sync_time', 'DATETIME2', '—', '—', '系统独有', '同步时间，同步时设置'],
    ]
)

add_heading('6.3 comMaterial（物料公司数据）补充字段', level=2)
add_para('comMaterial提供物料在公司层面的属性，可与comMaterialGroup配合使用：')
add_table(
    ['ERP字段(companyMaterial表)', '类型', '说明', '建议'],
    [
        ['MaterialId', 'String(40)', '物料代码（与comMaterialGroup一致）', '关联键'],
        ['MaterialTypeId', 'String(10)', '物料类型', '可覆盖comMaterialGroup的值'],
        ['MaterialCategoryId', 'String(20)', '物料类别', '需与ERP后台一致'],
        ['IsPurMat', 'Boolean', '是否采购物料', '可过滤只同步采购物料'],
        ['IsProdMat', 'Boolean', '是否生产物料', '可过滤只同步生产物料'],
        ['IsSalMat', 'Boolean', '是否销售物料', '可参考'],
        ['IsCalculateQty', 'Boolean', '是否核算库存数量', '可参考'],
        ['ValidityFromDate/ToDate', 'String(10)', '有效期', '可参考'],
    ]
)

add_heading('6.4 同步建议', level=2)
add_para('1. 使用comMaterialGroup的getlist接口拉取物料基础数据')
add_para('2. 可选使用comMaterial的getlist接口补充公司层面属性（需提供FOrgId公司代码）')
add_para('3. 使用data.condition过滤条件实现增量同步（如：MaterialId > \'上次最后一条\'）')
add_para('4. 利用hasnext和lastpkvalues实现分页遍历全部数据')
add_para('5. 同步时将source字段设为\'ERP\'，erp_sync_time设为当前时间')

doc.add_page_break()

# ===== 七、采购收货单同步字段映射 =====
add_heading('七、数据字段映射（采购收货单同步）', level=1)

add_heading('7.1 接口信息', level=2)
add_table(
    ['项目', '值'],
    [
        ['ERP接口(列表)', '/esb/erp/getlist.do'],
        ['ERP接口(明细)', '/esb/erp/get.do'],
        ['progid', 'purReceivingOrder_RC'],
        ['ERP表头', 'purReceivingOrderMaster'],
        ['ERP明细', 'purReceivingOrderDetail'],
        ['系统表头', 'requisition'],
        ['系统明细', 'requisition_item'],
        ['同步方向', 'ERP → 系统（单向拉取）'],
    ]
)

add_para('')
add_para('⚠️ 重要说明：', bold=True, color=(0xcc, 0x00, 0x00))
add_para('正航ERP API文档中提供的采购相关接口为"采购收货单(purReceivingOrder_RC)"，而非"请购单"或"采购订单"。')
add_para('系统的"采购订单管理"模块当前对应此接口。如ERP中有请购单/采购订单接口但未在文档中列出，需向正航确认。')

add_heading('7.2 表头字段映射（requisition ← purReceivingOrderMaster）', level=2)
add_table(
    ['系统字段(requisition表)', '类型', 'ERP字段(Master)', 'ERP类型', '映射说明', '备注'],
    [
        ['requisition_no', 'NVARCHAR(64)', 'BillNo', 'String(20)', '直接映射', '单据编号'],
        ['requisition_date', 'DATETIME2', 'BillDate', 'String(10)', '字符串→日期转换', '格式YYYY-MM-DD'],
        ['requester', 'NVARCHAR(64)', 'PersonId', 'String(20)', '代码→姓名转换', '需人员对照表'],
        ['department', 'NVARCHAR(128)', '—', '—', 'ERP无直接对应', '需通过PersonId查人员档案获取部门'],
        ['remark', 'NVARCHAR(512)', '—', '—', 'ERP无对应字段', '系统维护'],
        ['erp_sync_time', 'DATETIME2', '—', '—', '系统独有', '同步时间'],
        ['—', '—', 'TypeId', 'String(10)', '系统无对应字段', '单据类型，可新增字段'],
        ['—', '—', 'BizPartnerId', 'String(20)', '系统无对应字段', '供应商代码，可新增字段'],
        ['—', '—', 'IsPriceWithTax', 'Boolean', '系统无对应字段', '单价是否含税'],
    ]
)

add_heading('7.3 明细字段映射（requisition_item ← purReceivingOrderDetail）', level=2)
add_table(
    ['系统字段(requisition_item表)', '类型', 'ERP字段(Detail)', 'ERP类型', '映射说明', '备注'],
    [
        ['material_code', 'NVARCHAR(64)', 'MaterialId', 'String(40)', '直接映射', '物料代码'],
        ['material_name', 'NVARCHAR(256)', '—', '—', 'ERP明细无此字段', '需通过MaterialId查material表获取'],
        ['specification', 'NVARCHAR(512)', 'MaterialSpec', 'String(200)', '直接映射', '物料规格'],
        ['quantity', 'DECIMAL(18,4)', 'ReceivingSQty', 'Decimal(9)', '直接映射', '收货数量'],
        ['unit', 'NVARCHAR(16)', '—', '—', 'ERP明细无此字段', '需通过MaterialId查material表获取'],
        ['remark', 'NVARCHAR(512)', '—', '—', 'ERP无对应字段', '系统维护'],
        ['—', '—', 'RowCode', 'Int', '系统无对应字段', '行标识号'],
        ['—', '—', 'SPrice', 'Decimal(9)', '系统无对应字段', '交易单价，可新增字段'],
        ['—', '—', 'OAmount', 'Decimal(6)', '系统无对应字段', '金额，可新增字段'],
        ['—', '—', 'OAmountWithTax', 'Decimal(6)', '系统无对应字段', '含税金额，可新增字段'],
        ['—', '—', 'TaxId', 'String(20)', '系统无对应字段', '税代码'],
        ['—', '—', 'FromBillNo', 'String(20)', '系统无对应字段', '来源单号'],
    ]
)

add_heading('7.4 同步流程', level=2)
add_para('1. 调用getlist获取采购收货单列表（可按日期过滤：condition = "BillDate >= \'2026-07-01\'"）')
add_para('2. 对列表中的每一笔单据，调用getdata获取明细（pkvalues = BillNo）')
add_para('3. 明细中的material_name和unit需通过material_code关联本地material表补充')
add_para('4. 表头的department需通过PersonId关联comCompanyPerson人员档案获取')
add_para('5. 同步时设置erp_sync_time为当前时间')

doc.add_page_break()

# ===== 八、BOM同步字段映射 =====
add_heading('八、数据字段映射（BOM同步 — 待实现）', level=1)

add_para('当前系统尚未实现BOM管理功能，需新建数据库表和功能模块。以下为ERP BOM接口字段定义。')

add_heading('8.1 接口信息', level=2)
add_table(
    ['项目', '值'],
    [
        ['ERP接口(列表)', '/esb/erp/getlist.do'],
        ['ERP接口(明细)', '/esb/erp/get.do'],
        ['progid', 'ppBOM'],
        ['ERP表头', 'BOMMainInfo'],
        ['ERP子表', 'BOMSubMatInfo / BOMSubMatBatchQtyInfo / BOMSubMatInstallInfo / plsBOMSubReplacementDetail'],
    ]
)

add_heading('8.2 BOM表头字段（BOMMainInfo）', level=2)
add_table(
    ['ERP字段', '名称', '类型', '主键', '必填', '备注'],
    [
        ['BOMKeyId', '代码', 'String(50)', 'Y', 'Y', 'BOM唯一标识'],
        ['BOMKeyName', '名称', 'String(120)', '-', 'Y', 'BOM名称'],
        ['MaterialId', '母件', 'String(40)', '-', 'Y', '母件物料代码（关联material表）'],
        ['BizAttr', 'BOM属性', 'Byte', '-', 'Y', '0标准/1联产品/2结构/3虚拟/4通用/5订单'],
        ['Version', '版本号', 'String(20)', '-', 'Y', 'BOM版本'],
        ['BOMStyleId', 'BOM类型', 'String(10)', '-', 'Y', '固定值0001'],
        ['BOMTypeId', 'BOM用途类别', 'String(20)', '-', '-', '需同ERP后台'],
        ['BOMSerNo', '序', 'Int', '-', '-', '序号'],
        ['FromBizPartnerId', '客户代码', 'String(20)', '-', '-', ''],
        ['FromBillCategory', '来源单种类', 'Byte', '-', '-', ''],
        ['FromBillNo', '来源单号', 'String(20)', '-', '-', ''],
        ['FromRowCode', '来源标识号', 'Int', '-', '-', ''],
    ]
)

add_heading('8.3 BOM子件信息（BOMSubMatInfo）', level=2)
add_table(
    ['ERP字段', '名称', '类型', '主键', '必填', '备注'],
    [
        ['BOMKeyId', 'BOM代码', 'String(50)', 'Y', 'Y', '关联表头'],
        ['RowCode', '标识号', 'Int', 'Y', 'Y', '行号'],
        ['SubMaterialId', '子件代码', 'String(40)', '-', 'Y', '子件物料代码（关联material表）'],
        ['UnitQty', '用量', 'Decimal(9)', '-', 'Y', '子件用量'],
        ['BaseQty', '母件基数', 'Decimal(9)', '-', '-', ''],
        ['SubMatType', '子件来源', 'Byte', '-', '-', '0自备/1客户提供/2厂商提供/4文本'],
    ]
)

add_heading('8.4 建议的系统表结构', level=2)
add_para('新建BOM表：')
add_table(
    ['字段名', '类型', '说明'],
    [
        ['id', 'BIGINT IDENTITY', '主键'],
        ['bom_key_id', 'NVARCHAR(64)', 'BOM代码（ERP主键）'],
        ['bom_key_name', 'NVARCHAR(256)', 'BOM名称'],
        ['material_id', 'BIGINT', '母件物料ID（关联material表）'],
        ['material_code', 'NVARCHAR(64)', '母件物料代码'],
        ['biz_attr', 'TINYINT', 'BOM属性(0标准/1联产品/...)'],
        ['version', 'NVARCHAR(32)', '版本号'],
        ['erp_sync_time', 'DATETIME2', 'ERP同步时间'],
        ['created_at', 'DATETIME2', '创建时间'],
    ]
)

add_para('新建BOM明细表：')
add_table(
    ['字段名', '类型', '说明'],
    [
        ['id', 'BIGINT IDENTITY', '主键'],
        ['bom_id', 'BIGINT', '关联BOM表ID'],
        ['row_code', 'INT', '行标识号'],
        ['sub_material_id', 'BIGINT', '子件物料ID（关联material表）'],
        ['sub_material_code', 'NVARCHAR(64)', '子件物料代码'],
        ['unit_qty', 'DECIMAL(18,9)', '用量'],
        ['base_qty', 'DECIMAL(18,9)', '母件基数'],
        ['sub_mat_type', 'TINYINT', '子件来源(0自备/1客户/2厂商)'],
    ]
)

doc.add_page_break()

# ===== 九、生产订单同步 =====
add_heading('九、数据字段映射（生产订单同步 — 待实现）', level=1)

add_heading('9.1 接口信息', level=2)
add_table(
    ['项目', '值'],
    [
        ['ERP接口', '/esb/erp/getlist.do'],
        ['progid', 'ppProduceOrder'],
        ['ERP数据表', 'ppProduceOrder'],
        ['字段数量', '35个'],
    ]
)

add_heading('9.2 关键字段', level=2)
add_table(
    ['ERP字段', '名称', '类型', '备注'],
    [
        ['BillNo', '单据编号', 'String(20)', '主键'],
        ['BillDate', '单据日期', 'String(10)', 'YYYY-MM-DD'],
        ['TypeId', '单据类型', 'String(10)', ''],
        ['MaterialId', '母件', 'String(40)', '生产物料代码'],
        ['MaterialSpec', '物料规格', 'String(200)', ''],
        ['ProduceQty', '生产数量', 'Decimal(9)', ''],
        ['DemandBeginDate', '需求开始日期', 'String(10)', ''],
        ['DemandCompleteDate', '需求完工日期', 'String(10)', ''],
        ['DemandStockInDate', '需求入库日期', 'String(10)', ''],
        ['FromBOMKeyId', 'BOM', 'String(50)', '关联BOM'],
        ['ProduceState', '生产状态', 'Byte', '0未下达/1已下达/2部分完工/3已停工/4全部完工'],
        ['BizAttr', '业务属性', 'Int', '6标准/7重工/9拆解'],
        ['IssueTime', '下达时间', 'String(19)', ''],
        ['StockInQty', '入库数量', 'Decimal(9)', ''],
        ['Remark', '备注', 'String(2000)', ''],
    ]
)

doc.add_page_break()

# ===== 十、辅助接口 =====
add_heading('十、辅助接口（人员/部门/仓库）', level=1)

add_heading('10.1 人员档案（comCompanyPerson）', level=2)
add_para('用途：建立ERP人员代码(PersonId) ↔ 系统用户名(display_name)的对照关系')
add_para('特别用于：采购订单同步时，将ERP的PersonId转换为系统中的采购人员姓名')
add_table(
    ['项目', '值'],
    [
        ['接口', '/esb/erp/getlist.do'],
        ['progid', 'comCompanyPerson'],
    ]
)
add_para('建议在系统中新增"人员对照表"或在sys_user表新增erp_person_id字段，存储ERP人员代码。')

add_heading('10.2 部门（comDepartment）', level=2)
add_para('用途：获取ERP部门代码与名称的对照')
add_table(
    ['项目', '值'],
    [
        ['接口', '/esb/erp/getlist.do'],
        ['progid', 'comDepartment'],
    ]
)

add_heading('10.3 仓库（stkWareHouse）', level=2)
add_para('用途：如需扩展库存相关功能可使用')
add_table(
    ['项目', '值'],
    [
        ['接口', '/esb/erp/getlist.do'],
        ['progid', 'stkWareHouse'],
    ]
)

doc.add_page_break()

# ===== 十一、系统现状与缺口 =====
add_heading('十一、系统现状与缺口分析', level=1)

add_heading('11.1 已有功能', level=2)
add_table(
    ['功能模块', '状态', '数据库表', 'ERP对接状态'],
    [
        ['物料管理', '✅ 已实现', 'material', '有erp_sync_time字段，source字段，但未实现同步逻辑'],
        ['图纸管理', '✅ 已实现', 'drawing', '系统独有功能，无需ERP对接'],
        ['采购订单管理', '✅ 已实现', 'requisition + requisition_item', '有erp_sync_time字段，但未实现同步逻辑'],
        ['用户管理/权限', '✅ 已实现', 'sys_user', '无需ERP对接'],
        ['图纸分享', '✅ 已实现', 'share_token', '系统独有功能'],
        ['操作日志', '✅ 已实现', 'audit_log', '系统独有功能'],
    ]
)

add_heading('11.2 待开发功能', level=2)
add_table(
    ['功能模块', '状态', '需要的工作', 'ERP接口'],
    [
        ['ERP认证服务', '❌ 未实现', '新建EsbAuthService，实现token获取/缓存/自动刷新', 'auth'],
        ['物料同步服务', '❌ 未实现', '新建MaterialSyncService，定时/手动从ERP拉取物料', 'comMaterialGroup'],
        ['采购订单同步服务', '❌ 未实现', '新建RequisitionSyncService，定时从ERP拉取采购收货单', 'purReceivingOrder_RC'],
        ['人员对照', '❌ 未实现', '新增erp_person_id字段或人员对照表', 'comCompanyPerson'],
        ['BOM管理', '❌ 未实现', '新建BOM表/BOM明细表/Controller/Service/前端页面', 'ppBOM'],
        ['ERP配置管理', '❌ 未实现', '新建系统配置表存储ESB地址/appid/appsecret', '—'],
        ['同步日志', '❌ 未实现', '新建sync_log表记录每次同步的状态/数量/错误', '—'],
    ]
)

add_heading('11.3 需要新增的数据库表/字段', level=2)
add_para('1. 系统配置表（存储ESB地址/appid/appsecret等）')
add_para('2. 同步日志表（记录每次同步的执行情况）')
add_para('3. BOM主表 + BOM明细表（如需BOM功能）')
add_para('4. sys_user表新增 erp_person_id 字段（关联ERP人员代码）')
add_para('5. material表新增 material_category_id 字段（物料类别）')
add_para('6. requisition表新增 bill_type/supplier_id 字段（单据类型/供应商）')

doc.add_page_break()

# ===== 十二、实施建议 =====
add_heading('十二、实施建议与下一步', level=1)

add_heading('12.1 实施阶段划分', level=2)
add_table(
    ['阶段', '内容', '前置条件', '预计工作'],
    [
        ['第一阶段\n基础对接', 'ERP认证 + 物料同步', '获取ESB地址/appid/appsecret', 'EsbAuthService + MaterialSyncService + 配置管理'],
        ['第二阶段\n采购同步', '采购收货单同步 + 人员对照', '第一阶段完成 + 人员代码对照表', 'RequisitionSyncService + 人员对照'],
        ['第三阶段\nBOM管理', 'BOM同步 + BOM管理功能', '第二阶段完成', 'BOM表/Service/Controller/前端页面'],
        ['第四阶段\n生产订单', '生产订单同步', '第三阶段完成', '生产订单同步 + 前端展示'],
    ]
)

add_heading('12.2 第一步行动项（获取资料）', level=2)
add_para('请向正航软件实施人员或企业IT部门获取以下资料：', bold=True, color=(0xcc, 0x00, 0x00))
add_para('')
items = [
    '1. ESB服务地址及端口（如 http://192.168.1.100:8080）',
    '2. appid（应用ID，如 chiesb185a1fbdc22781f1）',
    '3. appsecret（应用密钥，如 2BB2D1D899C32BF091BEC8DDA9EBCB72）',
    '4. 确认ERP中是否有"请购单"或"采购订单"接口（文档中仅有"采购收货单"）',
    '5. 物料类型代码对照表（MaterialTypeId 的实际值清单）',
    '6. 物料类别代码对照表（MaterialCategoryId 的实际值清单）',
    '7. 计量单位代码对照表（UnitId 的实际值清单）',
    '8. 采购人员代码与姓名对照表（或ERP人员档案导出）',
    '9. 部门代码与名称对照表',
    '10. 公司代码（FOrgId，用于comMaterial查询）',
    '11. 确认同步频率要求（实时/定时）和数据范围（全量/增量）',
]
for item in items:
    add_para(item)

add_heading('12.3 技术建议', level=2)
add_para('1. 认证token缓存：获取后缓存，过期前5分钟自动刷新，避免频繁认证')
add_para('2. 增量同步：使用lastpkvalues分页 + condition条件过滤，避免全量拉取')
add_para('3. 错误重试：网络超时自动重试3次，记录失败日志')
add_para('4. 数据校验：同步前校验物料编码是否已存在，避免重复导入')
add_para('5. 手动触发+定时任务：支持管理员手动触发同步 + 定时自动同步')
add_para('6. 同步日志：记录每次同步的时间/数量/成功/失败/错误信息')
add_para('7. 回滚机制：同步失败时事务回滚，不更新已同步数据')

# ===== 保存 =====
output_path = r'e:\workspace\bom-drawing-system\正航T9_ERP对接需求文档.docx'
doc.save(output_path)
print(f'文档已保存: {output_path}')
